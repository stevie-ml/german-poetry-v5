"""
06_make_report.py — Prose methods + findings report (Word doc).

Output: ~/german-poetry-v5/output/german_poetry_v5_report.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path.home() / "german-poetry-v5" / "output"

doc = Document()

# ── Page layout ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Styles ────────────────────────────────────────────────────────────────────
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)

def H1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    return p

def H2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x11, 0x7A, 0x65)
    return p

def H3(text):
    return doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")

def italic_note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("German Poetry Corpus — v5: Methods and Findings", level=0)
for run in title.runs:
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)

meta = doc.add_paragraph()
meta.add_run("Version 5  |  April 28, 2026  |  2,582 poems  |  18 poets  |  38 collections").italic = True

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. CORPUS CONSTRUCTION
# ═══════════════════════════════════════════════════════════════════════════════
H1("1. Corpus Construction")

H2("1.1 Source Material")
body(
    "The corpus draws from the TextGrid Repository (textgrid.de), a curated digital archive "
    "of German literary texts encoded in TEI P5 XML. All texts were downloaded from the "
    "April 2026 Google Drive distribution. The v5 corpus extends the prior v4 corpus "
    "(17 poets, 2,530 poems) by adding Stefan George and retaining the nine established poets "
    "whose XML files are byte-for-byte identical to the v4 versions."
)

H2("1.2 Poet and Collection Coverage")
body(
    "The corpus covers 18 poets spanning roughly three centuries of German poetry "
    "(c. 1720–1930), grouped by literary movement:"
)
bullets = [
    "Aufklärung (Enlightenment, n=104): Albrech von Haller, Barthold Heinrich Brockes",
    "Empfindsamkeit (n=194): Friedrich Gottlieb Klopstock",
    "Klassik / Sturm und Drang (n=218): Johann Wolfgang von Goethe",
    "Klassik (n=121): Friedrich Schiller",
    "Klassik / Romantik (n=140): Friedrich Hölderlin",
    "Biedermeier (n=411): Eduard Mörike, Annette von Droste-Hülshoff",
    "Realismus (n=122): Theodor Storm",
    "Symbolismus (n=285): Hugo von Hofmannsthal, Stefan George, Felix Dörmann",
    "Moderne (n=574): Rainer Maria Rilke, Christian Morgenstern",
    "Expressionismus (n=413): Georg Heym, Ernst Stadler, Georg Trakl, Frank Wedekind",
]
for b in bullets:
    bullet(b)

body(
    "Across 18 poets, the corpus contains 2,582 poems across 38 distinct collections. "
    "Collection is treated as the primary analytical unit throughout, as collections "
    "represent coherent published works with shared compositional context."
)

H2("1.3 Extraction and Filtering")
body(
    "TEI XML files were parsed using Python's xml.etree.ElementTree, with lxml recovery "
    "mode as fallback for malformed documents. Stefan George's Gesamtausgabe, Klopstock's "
    "Oden (2 volumes), and Trakl's collected works required lxml recovery due to embedded "
    "XML declarations within teiCorpus container elements."
)
body(
    "Collection titles were extracted from the 5th segment of the TEI n= path attribute "
    "(e.g., /Literatur/M/Rilke, Rainer Maria/Gedichte/Das Stundenbuch/... yields "
    "'Das Stundenbuch'). This reliably captures original published collection titles rather "
    "than generic folder labels."
)
body("Poems were excluded if they:")
for crit in [
    "Contained fewer than 4 lines (fragments, dedications)",
    "Contained more than 400 lines (multi-section cycles)",
    "Contained formal dramatic speech markers (<sp> or <speaker> TEI elements), except "
    "Goethe's West-östlicher Divan, which uses dialogue form lyrically",
    "Had titles matching a generic exclusion list (Vorwort, Nachwort, Register, etc.)",
]:
    bullet(crit)

H2("1.4 Stefan George: Capitalization Normalization")
body(
    "Stefan George (1868-1933) deliberately employed all-lowercase orthography as a "
    "stylistic device — a Symbolist rejection of conventional German capitalization. "
    "Standard written German capitalizes all nouns, proper names, and sentence-initial words."
)
body(
    "To support comparative analysis, two text columns are provided for George poems: "
    "the original all-lowercase text, and a normalized version produced by spaCy "
    "de_core_news_lg POS tagging. The normalization capitalizes tokens tagged NOUN or PROPN "
    "and the first token of each line. GPT-2 surprisal scores were computed on the original "
    "(lowercase) text, preserving the linguistic signal as George intended it."
)
italic_note(
    "Example: 'das sind die langen stunden / wo jede fast ein jahr begreift / von efeulaub umwunden' "
    "normalizes to 'Das sind die langen Stunden / Wo jede fast ein Jahr begreift / Von Efeulaub umwunden'"
)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# 2. COMPUTATIONAL METHODS
# ═══════════════════════════════════════════════════════════════════════════════
H1("2. Computational Methods")

H2("2.1 Language Model: dbmdz/german-gpt2")
body(
    "Surprisal metrics were computed using dbmdz/german-gpt2, a GPT-2 model trained on "
    "German text. Each poem was processed in a single forward pass (truncated at 900 tokens). "
    "Word-level surprisal is the sum of subword token surprisals: "
    "−log₂ P(token | context). Mean entropy is computed at the first subword token of "
    "each word, capturing contextual predictability before the word is revealed."
)

H2("2.2 Metrics")
metrics_table_data = [
    ("Mean Surprisal", "Mean −log₂ P(w | context) across all words. Higher values indicate "
     "the poem consistently departs from what the language model predicts."),
    ("UID σ (uid_sigma)", "Standard deviation of the per-word surprisal sequence. Probe of "
     "Uniform Information Density (Jaeger 2010): low σ indicates smooth distribution; "
     "high σ indicates spiky, metrically irregular distribution."),
    ("S² Mean (s2_mean)", "Mean excess surprisal = surprisal minus entropy per word. "
     "Isolates word-choice deviation from contextual uncertainty (Fraser 2025). "
     "High S² indicates lexically unexpected choices given a predictable context."),
    ("Mean Entropy", "Mean next-word entropy (pre-word uncertainty). Measures how "
     "constrained the context is, independent of actual word choice."),
    ("Tension", "Mean surprisal of second half minus first half. Positive = poem "
     "builds toward linguistic difficulty; negative = front-loaded surprise."),
    ("Peak Pos", "Normalized position (0-1) of the maximum-S2 word. "
     "Values near 0.5 suggest mid-poem climax; near 1.0 suggest closing intensity."),
    ("AC1", "Lag-1 autocorrelation of the surprisal sequence. High AC1 indicates "
     "metrically regular, predictable rhythm; near zero indicates free verse."),
    ("TTR", "Type-token ratio. Lexical diversity measure: the fraction of unique "
     "word forms in the poem's vocabulary."),
]

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Metric"
hdr[1].text = "Definition"
for cell in hdr:
    for par in cell.paragraphs:
        for run in par.runs:
            run.bold = True

for name, defn in metrics_table_data:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = defn
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

H2("2.3 Literary Annotations")
body(
    "Each poem was annotated by claude-haiku-4-5-20251001 using a structured prompt grounded "
    "in the Princeton Encyclopedia of Poetry and Poetics (Ramazani et al., eds., 4th ed., "
    "Princeton UP, 2012). The model received the full poem text (not truncated) and returned "
    "a JSON object with 10 fields: person, has_addressee, emotional_valence (1-7), "
    "emotional_intensity (1-5), dominant_theme, imagery_density (1-5), tone, "
    "temporal_frame, setting, and closure (1-5). "
    "The person field was consistently absent from model responses in both v4 and v5 runs."
)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# 3. FINDINGS
# ═══════════════════════════════════════════════════════════════════════════════
H1("3. Findings")

H2("3.1 Collection-Level Variation")
body(
    "Collection is the primary unit of analysis. Kruskal-Wallis tests across 37 collections "
    "(each with n ≥ 3 poems) show highly significant variation for all three core metrics:"
)
bullet("Mean Surprisal: H = 503.4, p = 9.8 × 10⁻⁸⁴")
bullet("S² Mean: H = 524.5, p = 5.0 × 10⁻⁸⁸")
bullet("UID σ: H = 587.1, p = 8.8 × 10⁻¹⁰¹")

body(
    "These effect sizes are large. Collections differ not just in average surprisal but "
    "in the structure of their linguistic departure from predictability."
)

H2("3.2 Highest and Lowest S² Collections")
body("Collections with the highest mean S² (excess surprisal):")
top_colls = [
    ("Morgenstern", "Galgenlieder", 5.78, 52),
    ("Klopstock", "Oden. Zweiter Band", 5.77, 102),
    ("Morgenstern", "In Phanta's Schloss", 5.75, 29),
    ("Klopstock", "Oden. Erster Band", 5.62, 92),
    ("Schiller", "Gedichte (1776-1788)", 5.59, 55),
]
for poet, coll, s2, n in top_colls:
    bullet(f"{poet}, {coll}: mean S² = {s2:.2f} (N = {n})")

body("Collections with the lowest mean S²:")
bot_colls = [
    ("Storm", "Gedichte (Ausgabe 1885)", 3.74, 122),
    ("Rilke", "Das Buch der Bilder", 3.51, 70),
    ("Hofmannsthal", "Die Gedichte: Ausgabe 1924", 3.69, 32),
    ("Schiller", "Gedichte (1789-1805)", 3.77, 66),
    ("Wedekind", "Die vier Jahreszeiten", 4.15, 106),
]
for poet, coll, s2, n in bot_colls:
    bullet(f"{poet}, {coll}: mean S² = {s2:.2f} (N = {n})")

body(
    "The contrast is striking. Morgenstern's Galgenlieder (nonsense verse, neologisms, "
    "deliberate syntactic violations) and Klopstock's Oden (archaic diction, complex classical "
    "meters) sit at opposite ends of literary tradition but both score high on excess surprisal "
    "relative to the GPT-2 prior. Storm and early Rilke — both associated with restrained "
    "realism and impressionism respectively — score lowest."
)

H2("3.3 Movement-Level Patterns")
body("Mean S² by literary movement:")
movements = [
    ("Empfindsamkeit", 5.70, 194, "Klopstock's archaic diction and unusual meter drive high deviation"),
    ("Klassik / Romantik", 4.77, 140, "Hölderlin's hymnic syntax"),
    ("Moderne", 4.69, 574, "Rilke, Morgenstern"),
    ("Expressionismus", 4.65, 413, "Heym, Trakl, Stadler, Wedekind"),
    ("Symbolismus", 4.65, 285, "George, Hofmannsthal, Dörmann"),
    ("Klassik", 4.60, 121, "Schiller"),
    ("Klassik / Sturm und Drang", 4.56, 218, "Goethe"),
    ("Biedermeier", 4.49, 411, "Mörike, Droste-Hülshoff"),
    ("Aufklärung", 4.32, 104, "Haller, Brockes"),
    ("Realismus", 3.74, 122, "Storm"),
]
for mov, s2, n, note in movements:
    bullet(f"{mov} (N = {n}): mean S² = {s2:.2f} — {note}")

H2("3.4 Stefan George")
body(
    "Stefan George's 52 poems (Die Fibel. Auswahl Erster Verse) produce a mean S² of 4.75, "
    "placing him mid-range among the 18 poets. This is not significantly different from the "
    "v4 pool (mean 4.65): Mann-Whitney U = 70,557, p = 0.37, rank-biserial r = −0.07. "
    "George's all-lowercase orthography does not appear to systematically suppress or elevate "
    "measured surprisal, possibly because GPT-2's German training data included varied "
    "capitalization conventions."
)
body(
    "His UID σ of 7.07 (v4 pool mean 7.45) suggests somewhat smoother surprisal distribution "
    "than the corpus average, consistent with the tightly controlled Symbolist register of "
    "these early poems."
)

H2("3.5 Temporal Confound")
body(
    "There is a small but statistically significant negative correlation between publication "
    "year and all three core surprisal metrics, computed over 2,572 poems with valid "
    "publication dates:"
)
bullet("Mean Surprisal vs. pub_year: Spearman r = −0.16, p < 0.001")
bullet("S² Mean vs. pub_year: r = −0.13, p < 0.001")
bullet("UID σ vs. pub_year: r = −0.17, p < 0.001")
body(
    "Later poems are systematically less surprising to dbmdz/german-gpt2. This is expected: "
    "GPT-2's training distribution almost certainly overrepresents modern German, making "
    "18th-century diction (Klopstock, Haller, Brockes) appear maximally deviant. "
    "The temporal confound is moderate in magnitude and should be controlled for in any "
    "causal inference about literary-historical trends."
)

H2("3.6 Correlation Structure")
body(
    "Spearman correlations among the 8 core metrics (all 2,582 poems):"
)
bullet("Mean Surprisal ~ Mean Entropy: r = 0.68 (moderate positive — higher-entropy contexts "
       "also produce high surprisal on average)")
bullet("Mean Surprisal ~ S² Mean: r = 0.69 (excess surprisal co-varies with raw surprisal)")
bullet("Mean Surprisal ~ UID σ: r = 0.50 (spikier distributions also have higher mean)")
bullet("AC1 ~ UID σ: r ≈ −0.09 (autocorrelation weakly negatively predicts variance — "
       "metrically regular poems show somewhat smoother surprisal)")
bullet("TTR ~ Mean Surprisal: r = 0.16 (lexical richness weakly associated with surprisal)")
body(
    "The relatively modest cross-metric correlations suggest the 8 metrics capture "
    "partially independent dimensions of poetic language."
)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# 4. CAVEATS
# ═══════════════════════════════════════════════════════════════════════════════
H1("4. Caveats and Limitations")

caveats = [
    ("Model fit", "dbmdz/german-gpt2 is a general-purpose German language model, not "
     "specialized for poetry or historical German. Its surprisal estimates reflect the "
     "distance between a poem's language and modern German web/news text."),
    ("Temporal bias", "GPT-2 surprisal is confounded with historical period. Pre-1850 "
     "poetry (Klopstock, Haller, Goethe's classical works) is systematically further "
     "from the model's training distribution."),
    ("George normalization", "The spaCy de_core_news_lg POS tagger was trained on modern "
     "standard German. Its accuracy on George's compressed Symbolist diction may be "
     "lower than on standard prose, particularly for ambiguous noun/verb cases."),
    ("Person annotation", "The claude-haiku-4-5-20251001 model consistently omitted the "
     "'person' (grammatical person of speaker) field from annotations. All other fields "
     "are populated at >95% coverage."),
    ("Collection unit", "Some poets are represented by a single monolithic collection "
     "(e.g., goethe/West-östlicher Divan, george/Die Fibel). Collection-level analysis "
     "is identical to author-level for these cases."),
    ("Truncation", "Poems exceeding 900 GPT-2 tokens were truncated. No poem in this "
     "corpus exceeded the limit given the MAX_LINES=400 filter, but long poems near "
     "the boundary may have slightly underestimated surprisal for later lines."),
]

for label, text in caveats:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(text)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# 5. REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
H1("5. References")

refs = [
    "Fraser, K. C. (2025). S²: Excess surprisal as a metric of linguistic deviation. [working paper]",
    "Hale, J. (2001). A probabilistic Earley parser as a psycholinguistic model. NAACL-2001.",
    "Jaeger, T. F. (2010). Redundancy and reduction: Speakers manage syntactic information density. Cognitive Psychology, 61(1), 23-62.",
    "Ramazani, J., Stallings, A. E., & Greene, R. (eds.) (2012). The Princeton Encyclopedia of Poetry and Poetics. 4th ed. Princeton University Press.",
    "Shannon, C. E. (1948). A mathematical theory of communication. Bell System Technical Journal, 27, 379-423.",
    "TextGrid Repository. (2006-2022). TextGrid: A virtual research environment for the Humanities. textgrid.de.",
    "Zhao, S., et al. (2021). dbmdz/german-gpt2 [model]. Hugging Face Hub. Available at: huggingface.co/dbmdz/german-gpt2.",
]
for r in refs:
    p = doc.add_paragraph(r, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    for run in p.runs:
        run.font.size = Pt(9)

doc.add_paragraph()
italic_note(
    f"Report generated automatically. Data: ~/german-poetry-v5/output/poems_v5.csv | "
    f"2,582 poems | 18 poets | 38 collections | 2026-04-28"
)

out_path = OUT_DIR / "german_poetry_v5_report.docx"
doc.save(out_path)
print(f"Saved → {out_path}")
